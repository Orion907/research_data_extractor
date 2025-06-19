"""
Module for processing Word documents, specifically for extracting PICOTS tables. PICOTS
is an acronym for Population, Intervention, Comparator, Outcome, Timing, and Setting.
This module provides functions to extract PICOTS tables from Word documents and convert them
to structured format for use as a custom prompt in the research data extraction project.
"""
import logging
from docx import Document
from typing import List, Dict, Optional, Tuple
import re

# Configure logging
logger = logging.getLogger(__name__)

class DOCXProcessor:
    """
    Class for processing Word documents to extract PICOTS information from tables.
    """

     # Common PICOTS keywords to identify relevant tables
    PICOTS_KEYWORDS = {
        'P': ['population', 'participants', 'patients', 'subjects'],
        'I': ['intervention', 'exposure', 'treatment'],
        'C': ['comparison', 'control', 'comparator'],
        'O': ['outcomes', 'outcome', 'measures'],
        'T': ['timing', 'time', 'timeframe', 'duration'],
        'S': ['setting', 'study', 'design']
    }
    
    def __init__(self):
        logger.info("Initialized DOCX Processor for PICOTS extraction")
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """
        Extract all text from a Word document
        
        Args:
            docx_path (str): Path to the DOCX file
            
        Returns:
            str: Extracted text content
        """
        try:
            doc = Document(docx_path)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            full_text = '\n'.join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from {docx_path}")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Error extracting text from {docx_path}: {str(e)}")
            raise
    
    def extract_tables_from_docx(self, docx_path: str) -> List[List[List[str]]]:
        """
        Extract all tables from a Word document
        
        Args:
            docx_path (str): Path to the DOCX file
            
        Returns:
            List[List[List[str]]]: List of tables, each table is a list of rows, 
                                   each row is a list of cell contents
        """
        try:
            doc = Document(docx_path)
            tables_data = []
            
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        # Clean cell text
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    table_data.append(row_data)
                
                tables_data.append(table_data)
                logger.debug(f"Extracted table {table_idx + 1} with {len(table_data)} rows")
            
            logger.info(f"Extracted {len(tables_data)} tables from {docx_path}")
            return tables_data
            
        except Exception as e:
            logger.error(f"Error extracting tables from {docx_path}: {str(e)}")
            raise